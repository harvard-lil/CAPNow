import os
import re

import sys
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from docx import Document, RT
from docx.shared import Inches, Pt
from pyquery import PyQuery as pq
from copy import deepcopy

from scripts.utils import *
# # monkeypatch in missing namespace
# nsmap['vt'] = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"

template_path = 'sources/Case Template.docx'

def indent(elem, level=0):
    """
        Via http://effbot.org/zone/element-lib.htm
    """
    i = "\n" + level*"  "
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = i + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
        for elem in elem:
            indent(elem, level+1)
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = i

def save_xml(path, doc):
    doc_copy = deepcopy(doc.element)
    indent(doc_copy)
    #open(path+'.xml', 'wb').write(etree.tostring(doc_copy, pretty_print=True))

def load_doc(path):
    doc = Document(path)
    save_xml(path, doc)
    return doc, pq(doc.element, parser='xml')

def pq_copy(els):
    return pq(deepcopy(list(els)), namespaces=nsmap)

def has_text(el):
    return not re.match(r'^\s*$', pq(el).text())

def replace_element_contents(source_el, dest_el):
    pq(dest_el).empty().append(pq_copy(pq(source_el).children()))

def load_part(part):
    part_el = parse_xml(part.blob)
    part_pq = pq(part_el, parser='xml', namespaces=nsmap)
    return part, part_el, part_pq

def save_part(el, part):
    part._blob = etree.tostring(el)

def blank_run(run):
    """ Return blank version of run for style comparisons. """
    r1 = deepcopy(run)
    for attr in ('rsidR', 'rsidRPr'):
        if qn('w:'+attr) in r1.attrib:
            del r1.attrib[qn('w:'+attr)]
    pq(r1)('w|t').text('')
    return r1

def make_el(example_el, tag_name, attrs=None):
    if ':' in tag_name:
        tag_name = qn(tag_name)
    new_el = etree.Element(tag_name, nsmap=nsmap)
    if attrs:
        for k, v in attrs.items():
            new_el.attrib[qn(k)] = v
    return new_el

def parse_xml_fragment(example_el, xml):
    """ Parse xml using example_el for namespace information. """
    example_el = make_el(example_el, 'root')
    return list(parse_xml(etree.tostring(example_el, encoding=str)[:-2] + ">" + xml + "</root>"))

def split_run(old_run, texts):
    """
        Split a single run into multiple runs with supplied array of texts.
        Original run is removed, new runs are returned.
    """
    new_runs = []
    for text in texts:
        new_run = deepcopy(old_run)
        new_run.text = text
        new_runs.append(new_run)
    pq(old_run).after(pq(new_runs))
    remove_el(old_run)
    return new_runs

def convert(source_path, out_path, short_name, cite, year):

    ### TODO:
    # whitelist allowed tags
    # replace paragraph with .5 inch indented first line with a tab

    ### known changes:
    # tighter character spacing?
    # footnote numbers bold?
    # no space after footnote number?

    ### LOAD DATA ###

    # load docs
    source_doc, source_pq = load_doc(source_path)
    template_doc, template_pq = load_doc(template_path)

    # load footnotes
    footnotes_part, footnotes_el, footnotes_pq = load_part(source_doc.part.part_related_by(RT.FOOTNOTES))
    template_footnotes_part, template_footnotes_el, template_footnotes_pq = load_part(template_doc.part.part_related_by(RT.FOOTNOTES))

    ### COPY STYLES FROM TEMPLATE ###

    # copy styles, settings, and section formatting from template doc
    replace_element_contents(template_doc.styles._element, source_doc.styles._element)
    replace_element_contents(template_doc.settings._element, source_doc.settings._element)
    replace_element_contents(template_pq('w|sectPr')[0], source_pq('w|sectPr')[0])
    replace_element_contents(template_footnotes_pq('w|footnote').children()[0], footnotes_pq('w|footnote').children()[0])  # first footnote is the footnote separator

    ### HEADERS ###

    # delete existing header parts and copy in new header parts
    for rId, rel in list(source_doc.part.rels.items()):
        if rel.reltype == RT.HEADER:
            del source_doc.part.rels[rId]
    update_refs = {}
    header_parts = []
    for rId, rel in template_doc.part.rels.items():
        if rel.reltype == RT.HEADER:
            new_id = source_doc.part.rels._next_rId
            update_refs[rId] = new_id
            header_parts.append(load_part(rel.target_part))
            source_doc.part.rels.add_relationship(RT.HEADER, rel.target_part, new_id)
            source_doc.part.package.parts.append(rel.target_part)

    # update header references
    for header_ref in source_pq('w|headerReference'):
        header_ref.attrib[qn('r:id')] = update_refs[header_ref.attrib[qn('r:id')]]

    # fill in header values
    for header_part, header_el, header_pq in header_parts:
        header_pq("w|rStyle[w|val='HeaderYear']").closest('w|r')('w|t').text(year)
        header_pq("w|rStyle[w|val='HeaderCitation']").closest('w|r')('w|t').text(cite)
        short_name_par = Paragraph(header_pq("w|pStyle[w|val='HeaderCaseName']").closest('w|p')[0], None)
        short_name_par.clear()

        # italicize v. in party name
        if ' v. ' in short_name:
            party_a, party_b = short_name.split(' v. ', 2)
            short_name_par.add_run(party_a)
            vs_run = short_name_par.add_run(' v. ')
            vs_run.italic = True
            short_name_par.add_run(party_b)
        else:
            short_name_par.add_run(short_name)

    # set starting page number
    starting_page_number = cite.rsplit(' ',1)[-1]
    source_pq('w|sectPr').append(make_el(source_pq('w|sectPr')[0], 'w:pgNumType', {'w:start': starting_page_number}))

    ### TYPOGRAPHY ###

    # apply typography changes to body text and footnotes, adjusting variables that are different
    for query, allowed_styles, section_name, blockquote_style_name in (
            (source_pq, ('FootnoteReference',), 'body', 'Blockquote'),
            (footnotes_pq, ('FootnoteText', 'FootnoteSeparator', 'FootnoteReference'), 'footnote', 'FootnoteBlockquote')
    ):

        # clear existing styles
        ignore_removed_styles = ('NormalWeb',)
        for style_tag in query('w|pStyle,w|rStyle'):
            style_name = style_tag.attrib.get(qn('w:val'))
            if style_name not in allowed_styles:
                if style_name not in ignore_removed_styles:
                    print("Warning: removing unrecognized %s style %s." % (section_name, style_name))
                remove_el(style_tag)

        # mark block quotes
        for par in query('w|ind[w|left="720"]'):
            if qn('w:hanging') not in par.attrib:
                par = pq(par).closest('w|p')[0]
                par.style = blockquote_style_name

        # remove fonts and sizes
        remove_tags = ('sz', 'szCs', 'rFonts', 'ind', 'spacing', 'proofErr', 'bookmarkStart', 'bookmarkEnd', 'color[w|val="000000"]', 'lastRenderedPageBreak')
        for tag in remove_tags:
            query('w|'+tag).remove()

        # underline to italic
        for el in query('w|u'):
            if el.attrib.get(qn('w:val')) == 'double':
                el.tag = qn('w:smallCaps')
            else:
                el.tag = qn('w:i')
            el.attrib.clear()

        # combine consecutive runs with identical formatting
        query('w|t').attr(qn('xml:space'), 'preserve')  # add preserve to all t blocks for uniformity
        skip = 0
        for run in query('w|r'):

            # skip runs that have already been appended to previous run and detached
            if skip:
                skip -= 1
                continue

            blank_r1 = blank_run(run)
            while True:
                r2 = pq(run).next()
                if not r2:
                    break
                r2 = r2[0]
                if r2.tag != run.tag or etree.tostring(blank_r1) != etree.tostring(blank_run(r2)):
                    break
                run.text += r2.text
                remove_el(r2)
                skip += 1

        # text replacements
        for t in query('w|t'):
            text = t.text
            # fix dashes
            text = text.replace(" -- ", " — ")
            # remove double spaces
            text = re.sub(' +', ' ', text)
            # fix quotes
            for straight_quote, left_quote, right_quote in (('"', '“', '”'), ("'", '‘', '’')):
                if straight_quote not in text:
                    continue
                # right smart quotes
                text = re.sub(r'([a-zA-Z0-9.,?!;:\'\"])%s' % straight_quote, r'\1%s' % right_quote, text)
                text = re.sub(r'%s ' % straight_quote, r'%s ' % right_quote, text)
                # remaining are left smart quotes
                text = text.replace(straight_quote, left_quote)
            t.text = text

    ### FOOTNOTES ###

    footnote_tab = deepcopy(template_footnotes_pq('w|footnote:not([w|type]) w|r')[0])  # first run in template footnotes is a tab
    for footnote in footnotes_pq('w|footnote:not([w|type])'):

        # remove extra tabs from footnotes, add single tab
        for run in pq(footnote, namespaces=nsmap)('w|r'):
            if pq(run, namespaces=nsmap)('w|tab'):
                remove_el(run)
            else:
                pq(run).before(deepcopy(footnote_tab))
                break

        # make sure footnotes have FootnoteText style
        for par in pq(footnote, namespaces=nsmap)('w|p'):
            if not par.style:
                par.style = 'FootnoteText'

    ### CAPTION ###

    def skip_blanks(paragraphs, par_num):
        par_num += 1
        while not has_text(paragraphs[par_num]):
            par_num += 1
        return par_num

    # delete first four paragraphs
    pq(source_pq('w|p')[:4]).remove()

    paragraphs = source_pq('w|p')

    # format first paragraph
    par_num = 0
    paragraphs[par_num].style = 'CaseName'

    # process the case name so all-caps becomes small-caps:
    for run in pq(paragraphs[par_num])('w|r'):
        parts = re.split(r'([A-Z][A-Z]+)', run.text)
        if len(parts) > 1:
            new_runs = split_run(run, parts)
            for new_run in new_runs[1::2]:
                # every other part will be all-caps, so should become small-caps
                Run(new_run, None).font.small_caps = True
                new_run.text = new_run.text.title()

    par_num = skip_blanks(paragraphs, par_num)
    paragraphs[par_num].style = 'Dates'
    par_num = skip_blanks(paragraphs, par_num)
    paragraphs[par_num].style = 'Judges'
    par_num = skip_blanks(paragraphs, par_num)
    paragraphs[par_num].style = 'Categories'
    par_num = skip_blanks(paragraphs, par_num)

    while has_text(paragraphs[par_num]):
        paragraphs[par_num].style = 'Headnote'
        par_num += 2

    # extra space for last headnote
    Paragraph(paragraphs[par_num-2], None).paragraph_format.space_after = Pt(12)

    par_num = skip_blanks(paragraphs, par_num)
    while has_text(paragraphs[par_num]):
        paragraphs[par_num].style = 'History'
        par_num += 2

    par_num = skip_blanks(paragraphs, par_num)
    while has_text(paragraphs[par_num]):
        paragraphs[par_num].style = 'Appearance'
        par_num += 1

    # mark author name -- first sentence of first paragraph of case text
    par_num = skip_blanks(paragraphs, par_num)
    first_paragraph = Paragraph(paragraphs[par_num], source_doc._body)

    try:
        first_run = next(r for r in first_paragraph.runs if r.text.strip())
        first_run, second_run = split_run(first_run._element, first_run.text.split('.', 1))
        first_run.text = first_run.text.title() + "."
        Run(first_run, first_paragraph).style = "Author"
    except Exception as e:
        print("Warning: failed to detect author name. Searched this text: %s" % first_paragraph.text)
        raise

    # remove blank paragraphs
    # this has to come AFTER caption processing so we can tell sections apart
    for query in (source_pq, footnotes_pq('w|footnote:not([w|type])')):
        for p in query('w|p'):
            if not has_text(p):
                remove_el(p)

    ### HEADNOTE PAGE RANGES ###

    # replace highlighted headnote markers with bookmarks
    bookmarks = []
    for i, highlight_run in enumerate(source_pq("w|highlight[w|val='yellow']")):
        highlight_run = pq(highlight_run).closest('w|r')
        bookmark_name = "Headnote%s%s" % ("End" if i % 2 else "Start", int(i/2))
        highlight_run.after(pq([
            make_el(highlight_run[0], "w:bookmarkStart", {"w:id": str(i), "w:name": bookmark_name}),
            make_el(highlight_run[0], "w:bookmarkEnd", {"w:id": str(i)})]))
        remove_el(highlight_run[0])
        bookmarks.append(bookmark_name)

    # replace headnote page number references with bookmark shortcodes
    reference_template = """
        <w:fldSimple w:instr=" PAGEREF {bookmark_start} ">
            <w:r><w:rPr><w:noProof/></w:rPr><w:t>PRINT</w:t></w:r>
        </w:fldSimple>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> IF  </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> PAGEREF {bookmark_start} </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:rPr><w:noProof/></w:rPr><w:instrText>PRINT</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r><w:instrText xml:space="preserve"> = </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> PAGEREF {bookmark_end} </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:rPr><w:noProof/></w:rPr><w:instrText>PRINT</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r><w:instrText xml:space="preserve"> "" "-</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> PAGEREF {bookmark_end} </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:rPr><w:noProof/></w:rPr><w:instrText>PRINT</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r><w:instrText>"</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
    """
    for headnote in source_pq('w|pStyle[w|val="Headnote"]'):
        for run in pq(headnote).closest('w|p')('w|r'):
            run = pq(run)
            parts = re.split(r'\[.*?\]', run('w|t').text())
            if len(parts) > 1:
                new_els = []
                for i, part in enumerate(parts):
                    if i!= 0:
                        new_els.extend(parse_xml_fragment(run[0], reference_template.format(bookmark_start=bookmarks.pop(0), bookmark_end=bookmarks.pop(0))))
                    new_run = deepcopy(run[0])
                    pq(new_run)('w|t').text(
                        ("]" if i != 0 else "") + part + ("[" if i != len(parts) - 1 else "")
                    )
                    new_els.append(new_run)
                run.after(pq(new_els))
                remove_el(run[0])

    ### OUTPUT ###

    # write footnotes and headers
    save_part(footnotes_el, footnotes_part)
    for header_part, header_el, header_pq in header_parts:
        save_part(header_el, header_part)

    # save output
    #save_xml(out_path, source_doc)
    source_doc.save(out_path)




if __name__ == '__main__':
    # get paths
    in_path, out_path = sys.argv[1:]
    if os.path.isdir(out_path):
        out_path = os.path.join(out_path, os.path.basename(in_path))

    # get case info from source path
    name_abbreviation, cite, year = re.match(r'(.*), (\d+ Mass. \d+) \((\d{4})\)', in_path.rsplit('/', 1)[-1]).groups()

    convert(in_path, out_path, short_name, cite, year)